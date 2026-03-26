"""
File System and Document Tools with Security Controls

Safety features:
- Path validation against safe roots
- HTTPS-only URL validation
- File type restrictions
"""
import os
import logging
import aiofiles
import aiohttp
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional
from urllib.parse import urlparse
import hashlib
import magic
import json
import csv
from datetime import datetime
import io
import base64

# Document processing imports
import pdfplumber
from docx import Document
import httpx
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

from livekit.agents import function_tool, RunContext
from core.governance.policy import validate_safe_path, validate_download_url

logger = logging.getLogger(__name__)

# === DISCOVERY ===

@function_tool(name="list_directory")
async def list_directory(context: RunContext, path: str) -> str:
    """
    List files/folders at path. Returns: name, type (file/folder), size, last modified date.
    Expands ~ automatically.

    Args:
        path: Directory path to list
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return f"Access denied: path outside safe directories"

    if not os.path.isdir(full_path):
        return f"Not a directory: {full_path}"

    try:
        entries = []
        for entry in os.listdir(full_path):
            entry_path = os.path.join(full_path, entry)
            stat = os.stat(entry_path)

            entry_type = "file" if os.path.isfile(entry_path) else "folder"
            size = stat.st_size if entry_type == "file" else 0
            mtime = datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M')

            entries.append({
                "name": entry,
                "type": entry_type,
                "size": size,
                "modified": mtime
            })

        if not entries:
            return f"Directory is empty: {full_path}"

        result = f"Contents of {full_path}:\n\n"
        for entry in entries:
            result += f"• {entry['name']} ({entry['type']}) - {entry['size']} bytes - {entry['modified']}\n"

        return result

    except Exception as e:
        logger.error(f"Failed to list directory {full_path}: {e}")
        return f"Error listing directory: {e}"

@function_tool(name="search_files")
async def search_files(context: RunContext, pattern: str, path: str = "~") -> str:
    """
    Find files matching glob pattern under path. Example: search_files("*.pdf", "~/Documents")
    Returns matching paths with sizes. Max 50 results.

    Args:
        pattern: Glob pattern to match (e.g., "*.txt", "resume*")
        path: Root directory to search from (default: home)
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return f"Access denied: path outside safe directories"

    try:
        matches = []
        for match in Path(full_path).glob(f"**/{pattern}"):
            if match.is_file():
                matches.append({
                    "path": str(match),
                    "size": match.stat().st_size
                })

            if len(matches) >= 50:
                break

        if not matches:
            return f"No files matching '{pattern}' found in {full_path}"

        result = f"Found {len(matches)} files matching '{pattern}':\n\n"
        for match in matches[:20]:  # Limit output
            result += f"• {match['path']} ({match['size']} bytes)\n"

        if len(matches) > 20:
            result += f"\n... and {len(matches) - 20} more files (use more specific pattern)"

        return result

    except Exception as e:
        logger.error(f"Failed to search files: {e}")
        return f"Error searching files: {e}"

@function_tool(name="file_exists")
async def file_exists(context: RunContext, path: str) -> bool:
    """
    Check if path exists. Runs safe path validation first.

    Args:
        path: Path to check
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        logger.warning(f"File exists check denied for path outside safe roots: {full_path}")
        return False

    return os.path.exists(full_path)

@function_tool(name="file_metadata")
async def file_metadata(context: RunContext, path: str) -> Dict[str, Any]:
    """
    Return size, created, modified, extension, mime type.

    Args:
        path: Path to file
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return {"error": "Access denied: path outside safe directories"}

    if not os.path.isfile(full_path):
        return {"error": f"Not a file: {full_path}"}

    try:
        stat = os.stat(full_path)
        mime = magic.from_file(full_path, mime=True) if os.path.exists('/usr/bin/file') else "unknown"

        return {
            "path": full_path,
            "size": stat.st_size,
            "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "extension": Path(full_path).suffix,
            "mime_type": mime
        }
    except Exception as e:
        logger.error(f"Failed to get metadata for {full_path}: {e}")
        return {"error": f"Error getting metadata: {e}"}

@function_tool(name="file_hash")
async def file_hash(context: RunContext, path: str) -> str:
    """
    Return SHA256 hash of file. Used for: download verification, deduplication, change detection.

    Args:
        path: Path to file
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return f"Access denied: path outside safe directories"

    if not os.path.isfile(full_path):
        return f"Not a file: {full_path}"

    try:
        sha256_hash = hashlib.sha256()
        with open(full_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    except Exception as e:
        logger.error(f"Failed to hash file {full_path}: {e}")
        return f"Error computing hash: {e}"

@function_tool(name="count_file_lines")
async def count_file_lines(context: RunContext, path: str) -> int:
    """
    Return number of lines in a text file. Allows planner to compute read_file_chunk ranges before iterating.
    Example: count=800 → plan 4 chunks of 200 lines each.

    Args:
        path: Path to text file
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return -1

    if not os.path.isfile(full_path):
        return -1

    try:
        count = 0
        with open(full_path, 'r', encoding='utf-8') as f:
            for _ in f:
                count += 1
        return count
    except Exception as e:
        logger.error(f"Failed to count lines in {full_path}: {e}")
        return -1

# === READING ===

@function_tool(name="read_file")
async def read_file(context: RunContext, path: str) -> str:
    """
    Read file and return text content. Supported: .txt .md .py .json .csv .html .pdf (pdfplumber) .docx (python-docx)
    Truncates at 8000 tokens with truncation notice.

    Args:
        path: Path to file
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return f"Access denied: path outside safe directories"

    if not os.path.isfile(full_path):
        return f"Not a file: {full_path}"

    try:
        # Text files
        if full_path.endswith(('.txt', '.md', '.py', '.json', '.csv', '.html', '.css', '.js', '.yaml', '.yml', '.xml')):
            async with aiofiles.open(full_path, 'r', encoding='utf-8') as f:
                content = await f.read()

            # Simple token estimation (words / 0.75)
            tokens = len(content.split()) / 0.75
            if tokens > 8000:
                # Take first 6000 tokens worth of chars
                char_limit = int(6000 * 1.33)
                content = content[:char_limit]
                content += f"\n\n[TRUNCATED] File exceeds 8000 tokens limit. Showing first {int(char_limit / 1.33)} tokens."

            return content

        # PDF files
        elif full_path.endswith('.pdf'):
            text = ""
            with pdfplumber.open(full_path) as pdf:
                for page in pdf.pages[:20]:  # Limit to 20 pages
                    text += page.extract_text() or ""
                    text += "\n"

            tokens = len(text.split()) / 0.75
            if tokens > 8000:
                char_limit = int(6000 * 1.33)
                text = text[:char_limit]
                text += f"\n[TRUNCATED] PDF exceeds 8000 tokens limit."

            return text

        # Word documents
        elif full_path.endswith('.docx'):
            doc = Document(full_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            tokens = len(text.split()) / 0.75
            if tokens > 8000:
                char_limit = int(6000 * 1.33)
                text = text[:char_limit]
                text += f"\n[TRUNCATED] Document exceeds 8000 tokens limit."

            return text

        else:
            return f"Unsupported file type for read_file: {Path(full_path).suffix}"

    except Exception as e:
        logger.error(f"Failed to read file {full_path}: {e}")
        return f"Error reading file: {e}"

@function_tool(name="read_file_chunk")
async def read_file_chunk(context: RunContext, path: str, start_line: int, end_line: int) -> str:
    """
    Read specific line range from a file. Use count_file_lines first to plan chunk ranges.
    Example: read_file_chunk("report.txt", 1, 200)

    Args:
        path: Path to file
        start_line: Starting line number (1-indexed)
        end_line: Ending line number (inclusive)
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return f"Access denied: path outside safe directories"

    if not os.path.isfile(full_path):
        return f"Not a file: {full_path}"

    try:
        lines = []
        async with aiofiles.open(full_path, 'r', encoding='utf-8') as f:
            line_num = 1
            async for line in f:
                if start_line <= line_num <= end_line:
                    lines.append(line)
                line_num += 1
                if line_num > end_line:
                    break

        return "".join(lines)

    except Exception as e:
        logger.error(f"Failed to read file chunk {full_path}: {e}")
        return f"Error reading file chunk: {e}"

@function_tool(name="fetch_webpage")
async def fetch_webpage(context: RunContext, url: str) -> str:
    """
    Fetch visible text from URL via httpx + BeautifulSoup. HTTPS only. Strips nav/ads/scripts.
    Truncates to 6000 tokens.

    Args:
        url: HTTPS URL to fetch
    """
    # Validate URL
    valid, reason = validate_download_url(url)
    if not valid:
        return f"URL validation failed: {reason}"

    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=30) as response:
                if response.status != 200:
                    return f"HTTP error {response.status} fetching {url}"

                content = await response.text()

        # Parse with BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')

        # Remove navigation, ads, scripts
        for element in soup(['nav', 'header', 'footer', 'script', 'style', 'aside', 'noscript']):
            element.decompose()

        # Extract text
        text = soup.get_text(separator='\n', strip=True)

        # Clean up whitespace
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        cleaned = '\n'.join(lines)

        # Truncate
        tokens = len(cleaned.split()) / 0.75
        if tokens > 6000:
            char_limit = int(4500 * 1.33)
            cleaned = cleaned[:char_limit]
            cleaned += f"\n[TRUNCATED] Webpage exceeds 6000 tokens limit."

        return cleaned

    except Exception as e:
        logger.error(f"Failed to fetch webpage {url}: {e}")
        return f"Error fetching webpage: {e}"

# === FILE OPERATIONS ===

@function_tool(name="move_file")
async def move_file(context: RunContext, src: str, dst: str) -> str:
    """
    Move file. Validates both paths against safe roots.

    Args:
        src: Source file path
        dst: Destination file path
    """
    src_path = os.path.expanduser(src)
    dst_path = os.path.expanduser(dst)

    if not validate_safe_path(src_path):
        return f"Access denied: source path outside safe directories"
    if not validate_safe_path(dst_path):
        return f"Access denied: destination path outside safe directories"

    if not os.path.isfile(src_path):
        return f"Source does not exist or is not a file: {src_path}"

    try:
        os.makedirs(os.path.dirname(dst_path), exist_ok=True)
        os.rename(src_path, dst_path)
        return f"Moved file from {src_path} to {dst_path}"
    except Exception as e:
        logger.error(f"Failed to move file {src_path} to {dst_path}: {e}")
        return f"Error moving file: {e}"

@function_tool(name="copy_file")
async def copy_file(context: RunContext, src: str, dst: str) -> str:
    """
    Copy file. Validates both paths against safe roots.

    Args:
        src: Source file path
        dst: Destination file path
    """
    import shutil

    src_path = os.path.expanduser(src)
    dst_path = os.path.expanduser(dst)

    if not validate_safe_path(src_path):
        return f"Access denied: source path outside safe directories"
    if not validate_safe_path(dst_path):
        return f"Access denied: destination path outside safe directories"

    if not os.path.isfile(src_path):
        return f"Source does not exist or is not a file: {src_path}"

    try:
        os.makedirs(os.path.dirname(dst_path), exist_ok=True)
        shutil.copy2(src_path, dst_path)
        return f"Copied file from {src_path} to {dst_path}"
    except Exception as e:
        logger.error(f"Failed to copy file {src_path} to {dst_path}: {e}")
        return f"Error copying file: {e}"

@function_tool(name="download_file")
async def download_file(context: RunContext, url: str, destination: str) -> str:
    """
    Download URL to destination. Validates: HTTPS only, no blocked extensions, destination in safe roots.
    Returns: saved path, file size, SHA256 hash.

    Args:
        url: HTTPS URL to download from
        destination: Local path where file should be saved
    """
    # Validate URL
    valid, reason = validate_download_url(url)
    if not valid:
        return f"URL validation failed: {reason}"

    dst_path = os.path.expanduser(destination)

    if not validate_safe_path(dst_path):
        return f"Access denied: destination path outside safe directories"

    try:
        # Create directory
        os.makedirs(os.path.dirname(dst_path), exist_ok=True)

        # Download file
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=30) as response:
                if response.status != 200:
                    return f"HTTP error {response.status} downloading {url}"

                sha256_hash = hashlib.sha256()
                total_size = 0

                async with aiofiles.open(dst_path, 'wb') as f:
                    async for chunk in response.content.iter_chunked(8192):
                        await f.write(chunk)
                        sha256_hash.update(chunk)
                        total_size += len(chunk)

        return f"Downloaded {url} to {dst_path} ({total_size} bytes, SHA256: {sha256_hash.hexdigest()})"

    except Exception as e:
        logger.error(f"Failed to download {url}: {e}")
        return f"Error downloading file: {e}"

# === DOCUMENT CREATION ===

@function_tool(name="create_pdf")
async def create_pdf(context: RunContext, content: str, path: str) -> str:
    """
    Create PDF from markdown content using reportlab. Supports: h1/h2/h3, paragraphs, bullet lists.
    Returns saved path.

    Args:
        content: Markdown content to convert
        path: Output path (should end with .pdf)
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return f"Access denied: path outside safe directories"

    if not full_path.endswith('.pdf'):
        return f"Output path must end with .pdf"

    try:
        os.makedirs(os.path.dirname(full_path), exist_ok=True)

        doc = SimpleDocTemplate(full_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Simple markdown parsing
        for line in content.split('\n'):
            # Handle headers
            if line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Heading1']))
            # Handle bullet lists
            elif line.startswith('- ') or line.startswith('* '):
                story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            # Handle blank lines
            elif not line.strip():
                story.append(Spacer(1, 12))
            # Regular paragraphs
            else:
                story.append(Paragraph(line, styles['Normal']))

        doc.build(story)
        return f"Created PDF at {full_path}"

    except Exception as e:
        logger.error(f"Failed to create PDF at {full_path}: {e}")
        return f"Error creating PDF: {e}"

@function_tool(name="create_docx")
async def create_docx(context: RunContext, content: str, path: str) -> str:
    """
    Create Word document from markdown content using python-docx. Supports: h1/h2/h3, paragraphs, bullet lists.
    Returns saved path.

    Args:
        content: Markdown content to convert
        path: Output path (should end with .docx)
    """
    full_path = os.path.expanduser(path)

    if not validate_safe_path(full_path):
        return f"Access denied: path outside safe directories"

    if not full_path.endswith('.docx'):
        return f"Output path must end with .docx"

    try:
        os.makedirs(os.path.dirname(full_path), exist_ok=True)

        doc = Document()

        # Simple markdown parsing
        for line in content.split('\n'):
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                doc.add_paragraph(line)

        doc.save(full_path)
        return f"Created Word document at {full_path}"

    except Exception as e:
        logger.error(f"Failed to create Word document at {full_path}: {e}")
        return f"Error creating Word document: {e}"
